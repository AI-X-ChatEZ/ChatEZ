from typing import Union
from typing import List, Optional  # 추가된 코드
from datetime import datetime

from fastapi import FastAPI, UploadFile, File, Form, Body  # 추가된 코드
from fastapi.middleware.cors import CORSMiddleware
from openpyxl import load_workbook  # CORS보안 기능 설정
from pydantic import BaseModel
from urllib.parse import unquote
from opensearchpy import OpenSearch

from collections import deque
import tiktoken
import re
import os
import nltk
import docx
import chardet
import kss
import pandas as pd
import numpy as np
from konlpy.tag import Okt
from nltk.tag import pos_tag
from docx import Document
from rank_bm25 import BM25Okapi
from transformers import AutoTokenizer

from sklearn.preprocessing import MinMaxScaler
from sentence_transformers import SentenceTransformer, util
from unicodedata import normalize
import openai
from io import StringIO, BytesIO  # 추가된 코드
from dotenv import load_dotenv

load_dotenv() 

# 환경변수 설정
opensearch_host = os.getenv('OPENSEARCH_HOST')
opensearch_user = os.getenv('OPENSEARCH_USERNAME')
opensearch_password = os.getenv('OPENSEARCH_PASSWORD')
openai_api_key = os.getenv('OPENAI_API_KEY')

tokenizer = Okt()
nltk.download("punkt")
nltk.download("averaged_perceptron_tagger")

tokenizer_split = AutoTokenizer.from_pretrained("jhgan/ko-sroberta-multitask")

model = SentenceTransformer("jhgan/ko-sroberta-multitask")
model.max_seq_length = 512  #
tiktok = tiktoken.get_encoding("cl100k_base")
openai.api_key = openai_api_key

scaler = MinMaxScaler()


def preprocessing(text):
    text = str(text)
    text = normalize("NFKC", text)
    text = text.lower()
    text = re.sub(r"[^a-zA-Zㄱ-ㅣ가-힣0-9:%&#@\$€¥£.?!,~\-\n\(\) ]", " ", text)
    text = re.sub(" +", " ", text)
    text = re.sub("\n+", "\n", text)
    return text


def split_by_tokens(title, text, max_tokens=100):
    text = str(text)
    title = str(title)
    if not text:  # 값 비어있는 경우
        return []

    tokens = tokenizer_split.tokenize(text)
    if len(tokens) <= max_tokens:
        return ["title: " + title + ", content: " + text]
    splited_sentences = kss.split_sentences(text)
    result = []
    current_text = "title: " + title + ", content: "
    current_tokens = []
    for sentence in splited_sentences:
        sentence_tokens = tokenizer_split.tokenize(sentence)
        # 만약 문장 자체가 max_tokens 보다 크면, 해당 문장을 재귀적으로 분리
        while len(sentence_tokens) > max_tokens:
            sub_sentences = kss.split_sentences(sentence)
            if len(sub_sentences) <= 1:  # 더 이상 나눌 수 없는 경우
                break
            for sub_sentence in sub_sentences:
                sentence_tokens = tokenizer_split.tokenize(sub_sentence)
                if len(sentence_tokens) <= max_tokens:
                    result.append("title: " + title + ", content: " + sub_sentence)
                sentence = sentence.replace(sub_sentence, "").strip()
                sentence_tokens = tokenizer_split.tokenize(sentence)
        if len(current_tokens) + len(sentence_tokens) <= max_tokens:
            current_text += " " + sentence
            current_tokens.extend(sentence_tokens)
        else:
            if current_text:  # 이전 텍스트가 들어간게 있으면 이전 텍스만 result에 넣기
                result.append(current_text)
                current_text = "title: " + title + ", content: " + sentence
                current_tokens = sentence_tokens
            else:  # 이전 텍스트가 없으면 현재 텍스트만 result에 넣기
                result.append("title: " + title + ", content: " + sentence)
    if current_text != "title: " + title + ", content: ":
        result.append(current_text)
    return result


def mixed_tokenizer(text):
    # 한글과 영어를 분리
    korean_text = re.sub("[^가-힣\s]", "", text)
    english_text = re.sub("[^a-zA-Z\s]", "", text)
    korean_tokens = tokenizer.nouns(korean_text)
    english_tokens = pos_tag(nltk.word_tokenize(english_text))
    eng_nouns = [word for word, tag in english_tokens if tag.startswith("NN")]

    return korean_tokens + eng_nouns


from opensearchpy import OpenSearch

# AWS 인증 정보 설정
host = opensearch_host
port = 443
auth = (opensearch_user, opensearch_password)  # For testing only. Don't store credentials in code.

client = OpenSearch(
    hosts=[{"host": host, "port": port}],
    http_compress=True,  # enables gzip compression for request bodies
    http_auth=auth,
    # client_cert = client_cert_path,
    # client_key = client_key_path,
    use_ssl=True,
    verify_certs=True,
    ssl_assert_hostname=False,
    ssl_show_warn=False,
)

app = FastAPI()

class Upload(BaseModel):
    index: str
    files: List[UploadFile]  # 이건 더 알아보자


class Query(BaseModel):
    query: str
    history: Optional[List[str]] = None  # 메시지 히스토리를 위한 리스트


class DeleteFilesRequest(BaseModel):
    index: str
    file_ids: List[str]

# CORS 미들웨어 추가
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # "*" 모든 도메인의 연결을 허용
    allow_credentials=True,
    allow_methods=["*"],  # 허용된 HTTP 메소드
    allow_headers=["*"],  # 허용된 HTTP 헤더
)

def convert_bytes(num):
        for x in ["Bytes", "KB", "MB", "GB"]:
            if num < 1024.0:
                return f"{round(num)} {x}"
            num /= 1024.0

@app.get("/")
def read_root():
    return {"message": "deploy fastapi done"}

@app.post("/upload_files")
async def upload_files(index: str = Form(...), files: List[UploadFile] = File(...)):

    # Opensearch 서버 연결 확인
    if not client.ping():
        return "Opensearch 서버에 연결할 수 없습니다."

    #total_size 초기화
    total_size_in_bytes = 0


    for file in files:
        decoded_filename = unquote(file.filename)
        print(decoded_filename)
        file_content = await file.read()
        file_extension = decoded_filename.split(".")[-1]  # 파일 확장자 추출

        # 인코딩 확인
        bytes_io = BytesIO(file_content)
        rawdata = bytes_io.read()
        result = chardet.detect(rawdata)
        charenc = result["encoding"]
        total_size_in_bytes += len(file_content)

        # 파일 확장자에 따른 처리
        if file_extension == "docx":
            data = docx.Document(BytesIO(file_content))
            full_text = ""
            for paragraph in data.paragraphs:
                full_text += paragraph.text + "\n"
            full_text = preprocessing(full_text)

            data = split_by_tokens(decoded_filename, full_text)
            data = pd.DataFrame(data, columns=["data"])
            contents = data.explode("data").reset_index(drop=True)

        elif file_extension == "txt":
            if not charenc:
                content = rawdata.decode("utf-8")

            else:
                content = rawdata.decode(charenc)

            content = preprocessing(content)
            data = split_by_tokens(decoded_filename, content)
            data = pd.DataFrame(data, columns=["data"])
            contents = data.explode("data").reset_index(drop=True)

        # elif file_extension == "pdf":
        #     pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
        #     content = "\n".join([pdf_reader.getPage(i).extractText() for i in range(len(pdf_reader.pages))])

        elif file_extension in ["xlsx", "csv", "xls", "tsv"]:
            if file_extension in ["xlsx", "xls"]:
                data = pd.read_excel(BytesIO(file_content))

            elif file_extension in ["csv", "tsv"]:
                if not charenc:
                    charenc = "utf-8"

                if file_extension in ["csv"]:
                    data = pd.read_csv(BytesIO(file_content), encoding=charenc)
                    # data = pd.read_csv(BytesIO(file_content))
                else:
                    data = pd.read_csv(
                        BytesIO(file_content), sep="\t", encoding=charenc
                    )
                    # data = pd.read_csv(BytesIO(file_content), sep="\t")

            if data.isna().any().any():
                data.fillna(" ", inplace=True)

            data.applymap(preprocessing)
            columns_len = len(data.columns)

            if columns_len <= 2:
                if (
                    (columns_len == 2)
                    and ("Question" in data.columns)
                    and ("Answer" in data.columns)
                ):
                    data = data.apply(
                        lambda row: split_by_tokens(row.Question, row.Answer), axis=1
                    )
                elif columns_len == 2:
                    if (
                        data[
                            data[data.columns[0]].str.contains("?", regex=False)
                        ].shape[0]
                        > data[
                            data[data.columns[1]].str.contains("?", regex=False)
                        ].shape[0]
                    ) or (
                        data[data.columns[0]].str.len().sum()
                        < data[data.columns[1]].str.len().sum()
                    ):
                        data.rename(columns={data.columns[0]: "Question"}, inplace=True)
                        data.rename(columns={data.columns[1]: "Answer"}, inplace=True)
                    else:
                        data.rename(columns={data.columns[1]: "Question"}, inplace=True)
                        data.rename(columns={data.columns[0]: "Answer"}, inplace=True)

                    data = data.apply(
                        lambda row: split_by_tokens(row.Question, row.Answer), axis=1
                    )

                elif columns_len == 1:
                    data.rename(columns={data.columns[0]: "context"}, inplace=True)
                    data = data.apply(
                        lambda row: split_by_tokens(decoded_filename, row.context), axis=1
                    )

                data = pd.DataFrame(data, columns=["data"])
                contents = data.explode("data").reset_index(drop=True)

            else:
                data["data"] = ", ".join(
                    [f"{col}:{data[col].iloc[0]}" for col in data.columns]
                )
                data = data.apply(
                    lambda row: split_by_tokens(decoded_filename, row.data), axis=1
                )
                data = pd.DataFrame(data, columns=["data"])
                contents = data.explode("data").reset_index(drop=True)

        else:
            contents = {'data' : "지원하지 않는 확장자"}

        current_time = datetime.now()
        formatted_time = current_time.strftime("%H:%M:%S.%f")[:-4]
        upload_time = current_time.strftime("%Y-%m-%d")

        for i in range(len(contents)):
            content = contents["data"][i]
            new_index = index + str(i)
            size_in_bytes = len(content)
            readable_size = convert_bytes(size_in_bytes)
            readable_total_size = convert_bytes(total_size_in_bytes)

            BM25_tokenized = mixed_tokenizer(content)
            # print(type(BM25_tokenized))
            SBERT_Embedding = model.encode(content).tolist()
            # print(type(SBERT_Embedding))
            file_name = decoded_filename.rsplit('.', 1)[0]
            document = {
                "documentId": f"{decoded_filename}_{formatted_time}",
                "size": readable_size,
                "totalSize": readable_total_size,
                "contents": content.strip(),
                "name": file_name,  # 파일 이름만 저장
                "contentType": file_extension,
                "uploadTime": upload_time,
                "BM25_tokenized": BM25_tokenized,
                "SBERT_Embedding": SBERT_Embedding,
            }

            unique_key = f"{decoded_filename}_{formatted_time}_{i}"
            response = client.index(index=index, id=unique_key, body=document)
            # print(response)

    return "Opensearch에 업로드 성공"

@app.post("/update_files")
async def upload_files(index: str = Form(...), files: List[UploadFile] = File(...)):

    # Opensearch 서버 연결 확인
    if not client.ping():
        return "Opensearch 서버에 연결할 수 없습니다."

    #total_size 초기화
    total_size_in_bytes = 0

    for file in files:
        decoded_filename = unquote(file.filename)
        file_content = await file.read()
        file_extension = decoded_filename.split(".")[-1]  # 파일 확장자 추출

        # 인코딩 확인
        bytes_io = BytesIO(file_content)
        rawdata = bytes_io.read()
        result = chardet.detect(rawdata)
        charenc = result["encoding"]
        total_size_in_bytes += len(file_content)

        # 파일 확장자에 따른 처리
        if file_extension == "docx":
            data = docx.Document(BytesIO(file_content))
            full_text = ""
            for paragraph in data.paragraphs:
                full_text += paragraph.text + "\n"
            full_text = preprocessing(full_text)

            data = split_by_tokens(decoded_filename, full_text)
            data = pd.DataFrame(data, columns=["data"])
            contents = data.explode("data").reset_index(drop=True)

        elif file_extension == "txt":
            if not charenc:
                with open(BytesIO(file_content), "rb") as file:
                    content = file.read()
                    content = content.decode("utf-8")

            else:
                with BytesIO(file_content) as file_encoding:
                    rawdata = file_encoding.read()
                    content = rawdata.decode(charenc)

            content = preprocessing(content)
            data = split_by_tokens(decoded_filename, content)
            data = pd.DataFrame(data, columns=["data"])
            contents = data.explode("data").reset_index(drop=True)

        # elif file_extension == "pdf":
        #     pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
        #     content = "\n".join([pdf_reader.getPage(i).extractText() for i in range(len(pdf_reader.pages))])

        elif file_extension in ["xlsx", "csv", "xls", "tsv"]:
            if file_extension in ["xlsx", "xls"]:
                data = pd.read_excel(BytesIO(file_content))

            elif file_extension in ["csv", "tsv"]:
                if not charenc:
                    charenc = "utf-8"

                if file_extension in ["csv"]:
                    data = pd.read_csv(BytesIO(file_content), encoding=charenc)
                    # data = pd.read_csv(BytesIO(file_content))
                else:
                    data = pd.read_csv(
                        BytesIO(file_content), sep="\t", encoding=charenc
                    )
                    # data = pd.read_csv(BytesIO(file_content), sep="\t")

            if data.isna().any().any():
                data.fillna(" ", inplace=True)

            data.applymap(preprocessing)
            columns_len = len(data.columns)

            if columns_len <= 2:
                if (
                    (columns_len == 2)
                    and ("Question" in data.columns)
                    and ("Answer" in data.columns)
                ):
                    data = data.apply(
                        lambda row: split_by_tokens(row.Question, row.Answer), axis=1
                    )
                elif columns_len == 2:
                    if (
                        data[
                            data[data.columns[0]].str.contains("?", regex=False)
                        ].shape[0]
                        > data[
                            data[data.columns[1]].str.contains("?", regex=False)
                        ].shape[0]
                    ) or (
                        data[data.columns[0]].str.len().sum()
                        < data[data.columns[1]].str.len().sum()
                    ):
                        data.rename(columns={data.columns[0]: "Question"}, inplace=True)
                        data.rename(columns={data.columns[1]: "Answer"}, inplace=True)
                    else:
                        data.rename(columns={data.columns[1]: "Question"}, inplace=True)
                        data.rename(columns={data.columns[0]: "Answer"}, inplace=True)

                    data = data.apply(
                        lambda row: split_by_tokens(row.Question, row.Answer), axis=1
                    )

                elif columns_len == 1:
                    data.rename(columns={data.columns[0]: "context"}, inplace=True)
                    data = data.apply(
                        lambda row: split_by_tokens(decoded_filename, row.context), axis=1
                    )

                data = pd.DataFrame(data, columns=["data"])
                contents = data.explode("data").reset_index(drop=True)

            else:
                data["data"] = ", ".join(
                    [f"{col}:{data[col].iloc[0]}" for col in data.columns]
                )
                data = data.apply(
                    lambda row: split_by_tokens(decoded_filename, row.data), axis=1
                )
                data = pd.DataFrame(data, columns=["data"])
                contents = data.explode("data").reset_index(drop=True)

        else:
            contents = {'data' : "지원하지 않는 확장자"}

        current_time = datetime.now()
        formatted_time = current_time.strftime("%H:%M:%S.%f")[:-3]
        upload_time = current_time.strftime("%Y-%m-%d")

        for i in range(len(contents)):
            content = contents["data"][i]
            new_index = index + str(i)
            size_in_bytes = len(content)
            readable_size = convert_bytes(size_in_bytes)
            readable_total_size = convert_bytes(total_size_in_bytes)


            BM25_tokenized = mixed_tokenizer(content)
            # print(type(BM25_tokenized))
            SBERT_Embedding = model.encode(content).tolist()
            # print(type(SBERT_Embedding))
            file_name = decoded_filename.rsplit('.', 1)[0]
            
            document = {
                "documentId": f"{decoded_filename}_{formatted_time}",
                "size": readable_size,
                "totalSize": readable_total_size,
                "contents": content.strip(),
                "name": file_name,  # 파일 이름만 저장
                "contentType": file_extension,
                "uploadTime": upload_time,
                "BM25_tokenized": BM25_tokenized,
                "SBERT_Embedding": SBERT_Embedding,
            }

            response = client.index(
                index=index,  # 이것은 문서 ID가 아닌 OpenSearch 인덱스 이름이어야 합니다
                body=document
            )
            print(response)

    return "Opensearch에 업로드 성공"

@app.post("/delete_files")
async def delete_files(request: DeleteFilesRequest):
    # Opensearch 서버 연결 확인
    if not client.ping():
        return "Opensearch 서버에 연결할 수 없습니다."

    for file_id in request.file_ids:

        print(file_id)

        # _id 필드 안에서 해당 문자열을 포함하는 모든 문서 검색
        search_query = {
            "query": {
                "term": {
                    "documentId.keyword": file_id
                }
            }
        }

        response = client.search(index=request.index, body=search_query)
        #print(response)
        # 검색된 문서들 삭제
        for hit in response["hits"]["hits"]:
            client.delete(index=request.index, id=hit["_id"])
            print(f"Deleted document with ID: {hit['_id']}")

    return "Opensearch에서 파일 삭제 성공"

@app.post("/handle_query/{index}")
def handle_query(index: str, query: Query):
    # print(index)
    message = query.query  # 메시지 내용을 가져옴
    message_history = query.history  # 메시지 히스토리를 가져옴
    
    # print("Current Message:", message)  # 콘솔에 메시지 내용 출력
    # print("Message History:", message_history)
    preconvs= deque(message_history) 
    print("preconvs deque: ", preconvs)

    search_query = {
    "query": {
        "match_all":{}
        }
    } 
    # response = client.search(index=index, body=search_query)
    fields = ['contents', "BM25_tokenized", "SBERT_Embedding"]
    response = client.search(index=index, _source = fields)

    data  =[]
    for hit in response["hits"]["hits"]:
        data.append(hit["_source"]) 
    
    df = pd.DataFrame(data)

    query_token = mixed_tokenizer(message)
    # query_token = tokenizer.tokenize(message)
    bm25 = BM25Okapi(df["BM25_tokenized"])
    df["BM25_score"] = list(bm25.get_scores(query_token))
    df["BM25_score"] = scaler.fit_transform(df[["BM25_score"]])

    query_embedding = model.encode(message)
    df["cossim_score"] = df["SBERT_Embedding"].apply(
        lambda x: util.cos_sim(x, query_embedding)
    )
    df["cossim_score"] = scaler.fit_transform(df[["cossim_score"]])

    df["Hybrid_score"] = (
        df["BM25_score"] * 4 + df["cossim_score"] * 6
    )
    df = df.sort_values(by=["Hybrid_score"], ascending=False)
    selected_docs = list(df.iloc[:5]["contents"])

    while True:
        try:
            # preconv = ' '.join(preconvs)
            preconv = ""

            if preconvs:
                for i in preconvs:
                    preconv += " question : " + i[0]

            selected_doc = " ".join(selected_docs)
            print(selected_docs)

            system_content = "You are a helpful, respectrul, friendly chatbot."

            instruction = """knowledge:
            1. If someone asks for the price of 3 items, multiply the price of the item by 3 to find the total price.
            2. Find the price for free delivery. If there is information about a delivery fee, remember it, but do not answer it yet.
            3. If the total price does not meet the price for free delivery, add the delivery fee to the total price, if there is one. This is the final price.
            instruction:
            Given a context, identify and extract the sentences or parts that are most relevant to the user's query. Ignore unrelated sections and ensure that the response is natural and free of hallucinations. If multiple product names are mentioned in the extracted document, ask for clarification on which product is being referred to.

            Answer in korean.
            """
            prev_conversation = f"Previous conversation:\n{preconv}"
            context = f"Context:\n{selected_doc}"
            user_query = f"User query:\n{message}"
            user_content = prev_conversation + context + user_query + instruction      

            message=[ {"role": "system",
                        "content": system_content},
                        {"role": "user",
                        "content": user_content}]

            num_tokens = 6  # start, role/content*2, end

            for i in range(len(message)):
                for value in message[i].values():
                    num_tokens += len(tiktok.encode(value))

            if num_tokens >= 3697:  # 4097 - 400(max_tokens)
                if preconv:
                    print("너무 길엉 이전 대화 없애쟝")
                    preconvs.popleft()
                    continue
                else:
                    return "질문이 너무 깁니다. 다시 질문해주세요."
                    break

            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=message,
                temperature=0,
                max_tokens=400,
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0,
                request_timeout=30,
            )
            answer = response["choices"][0]["message"]["content"]

            return answer
            break

        except Exception as e:
            print(e)
            if "authentication" in str(e).lower():
                return "인증 에러: ChatEZ 관리자에게 문의해주세요"

            elif "Quota exceeded" in str(e):
                return "API 사용량을 초과했습니다. ChatEZ 관리자에게 문의해주세요."

            elif "Model not available" in str(e):
                return "요청한 모델을 사용할 수 없습니다. ChatEZ 관리자에게 문의해주세요"

            elif "Invalid input" in str(e):
                return "입력값이 잘못되었습니다. 입력 파라미터를 확인해주세요."

            else:
                # 이 외의 에러 처리 로직
                return f"알 수 없는 에러 발생, 다시 시도해주세요."

            break